from __future__ import annotations

import hashlib
import tempfile
from pathlib import Path
from typing import List, Optional

from src.tools.document.parsers.types import DocumentContent, ImageContent, TableContent

try:
    import docx  # type: ignore
    from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
except Exception:
    docx = None  # type: ignore
    RT = None  # type: ignore


class DocxParser:
    """DOCX 文档解析器。

    支持提取：
    - 段落文本
    - 表格内容（结构化）
    - 嵌入图片
    """

    def __init__(self, output_dir: Optional[Path] = None) -> None:
        """初始化解析器。

        Args:
            output_dir: 图片输出目录，默认使用临时目录
        """
        self._output_dir = output_dir

    def parse(self, path: Path) -> DocumentContent:
        """解析 DOCX 文档。

        Args:
            path: DOCX 文件路径

        Returns:
            结构化的文档内容
        """
        if docx is None:
            return DocumentContent(
                text="",
                metadata={
                    "path": str(path),
                    "format": "docx",
                    "warning": "python-docx not installed.",
                },
            )

        try:
            document = docx.Document(str(path))

            # 提取文本段落
            paragraphs = self._extract_paragraphs(document)

            # 提取表格
            tables = self._extract_tables(document)
            table_texts = [self._table_to_text(t) for t in tables]

            # 提取图片
            images = self._extract_images(document, path)

            metadata = {
                "path": str(path),
                "format": "docx",
                "parser": "python-docx",
                "paragraphs": len(paragraphs),
                "tables": len(tables),
                "images": len(images),
            }

            # 合并文本内容
            text = "\n".join(paragraphs)
            if table_texts:
                text += "\n\n" + "\n\n".join(table_texts)

            return DocumentContent(
                text=text,
                images=images,
                tables=tables,
                metadata=metadata,
            )
        except Exception as exc:
            return DocumentContent(
                text="",
                metadata={
                    "path": str(path),
                    "format": "docx",
                    "warning": f"DOCX parsing failed: {exc}",
                },
            )

    @staticmethod
    def _extract_paragraphs(document) -> List[str]:
        """提取文档段落。"""
        return [para.text for para in document.paragraphs if para.text.strip()]

    @staticmethod
    def _extract_tables(document) -> List[TableContent]:
        """提取文档中的表格。"""
        tables: List[TableContent] = []
        for idx, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append(
                TableContent(
                    raw=rows,
                    description=f"表格 {idx}",
                )
            )
        return tables

    @staticmethod
    def _table_to_text(table: TableContent) -> str:
        """将表格转换为文本表示。"""
        if not table.raw:
            return ""
        lines = []
        for row in table.raw:
            lines.append("\t".join(str(cell) for cell in row))
        return "\n".join(lines)

    def _extract_images(self, document, doc_path: Path) -> List[ImageContent]:
        """提取文档中的嵌入图片。"""
        images: List[ImageContent] = []

        if RT is None:
            return images

        # 确定图片输出目录（优先使用配置的目录，否则使用文档同级 images 目录）
        output_dir = self._output_dir
        if output_dir is None:
            # 默认保存到文档同级的 images 目录，而不是临时目录
            output_dir = doc_path.parent / "images" / doc_path.stem
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            # 遍历文档中的所有关系，找到图片
            for rel_id, rel in document.part.rels.items():
                if rel.reltype == RT.IMAGE:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob

                        # 生成唯一文件名
                        content_hash = hashlib.md5(image_bytes).hexdigest()[:8]
                        ext = self._get_image_extension(image_part.content_type)
                        filename = f"{doc_path.stem}_{rel_id}_{content_hash}{ext}"
                        image_path = output_dir / filename

                        # 保存图片
                        image_path.write_bytes(image_bytes)

                        images.append(
                            ImageContent(
                                path=str(image_path),
                                caption="",
                                description="",
                                image_type="embedded",
                            )
                        )
                    except Exception:
                        # 单个图片提取失败不影响其他图片
                        continue
        except Exception:
            # 图片提取失败不影响文档解析
            pass

        return images

    @staticmethod
    def _get_image_extension(content_type: str) -> str:
        """根据 MIME 类型获取文件扩展名。"""
        type_map = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/tiff": ".tiff",
            "image/webp": ".webp",
        }
        return type_map.get(content_type, ".png")
